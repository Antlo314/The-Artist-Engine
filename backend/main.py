import asyncio
import os
import shutil
import subprocess
import uuid
import json
import time
import sys
import io
import PyPDF2
from docx import Document
from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Request, BackgroundTasks, Depends
from fastapi.responses import FileResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from pydub import AudioSegment
import traceback
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import Optional
from dotenv import load_dotenv

# Load Environment Variables from local .env
load_dotenv(dotenv_path=".env")
# Also load repo-root .env when running from backend/
load_dotenv(dotenv_path=os.path.join(os.path.dirname(__file__), "..", ".env"))

from agent_memory import AgentMemory
from gig_sources import ticketmaster_available, fetch_ticketmaster_venues
from bandsintown_client import bandsintown_available, venues_near_city as bit_venues_near_city
from musicbrainz_client import search_artist, artist_releases, epk_bundle
from legal_linter import lint_text, DISCLAIMER as LEGAL_DISCLAIMER
from audio_meters import measure_audio
from stem_engine import separate_stems
from crm import (
    init_crm_tables,
    list_leads,
    upsert_lead,
    bulk_upsert_leads,
    move_lead,
    delete_lead,
    add_pitch,
    list_pitches,
    add_activity,
    list_activity,
    export_all,
    import_all,
    create_reset_token,
    reset_password_with_code,
    count_leads,
    list_tasks,
    add_task,
    set_task_done,
    PIPELINE_STAGES,
)
from free_data import (
    free_stack_manifest,
    GENRE_PRESETS,
    RELEASE_CHECKLIST,
    tour_route,
    discogs_search,
    nominatim_geocode,
    infer_country,
)
from founding_auth import (
    AppUser,
    auth_configured,
    auth_required,
    require_founding_user,
    assert_quota,
    record_usage,
    get_usage_snapshot,
    acquire_master_slot,
    release_master_slot,
    DAILY_LIMITS,
    init_db,
    register_user,
    login_user,
    login_or_register_google,
    logout_token,
    list_users_admin,
)
from entitlements import list_plans_public, CREDIT_PACKS, feature_allowed, credit_cost
from billing import (
    redeem_promo,
    create_checkout_session,
    handle_stripe_webhook,
    get_credits_snapshot,
    promo_stats,
    admin_adjust_credits,
    admin_set_plan,
    admin_add_note,
    effective_plan_for_user,
    ensure_monthly_grant,
    active_multiplier,
)
from user_profiles import (
    ensure_profile,
    get_profile,
    save_profile,
    recompute_stats_from_usage,
    seed_demo_artists,
    init_profile_tables,
)
from supabase_auth import (
    supabase_configured,
    fetch_supabase_user,
    optional_sync_app_user_row,
)

# Google GenAI Integration
try:
    from google import genai
    from google.genai import types
except ImportError:
    print("[SYSTEM WARNING] google-genai library missing. Install via pip install google-genai")

# ---------------------------------------------------------------------------
# OMEGA-TIER SYSTEM CONFIGURATION
# ---------------------------------------------------------------------------

app = FastAPI(
    title="The Source Engine API",
    description="Backend for The Source Engine — booking, contracts, mastering, and CRM tools for independent artists.",
    version="3.0.0"
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:5173",
        "http://127.0.0.1:5173",
        "https://the-artist-engine-lime.vercel.app",
        "https://artist-engine.vercel.app",
        "https://theartistengine.com",
        "https://www.theartistengine.com",
        "https://thesourceengine.com",
        "https://www.thesourceengine.com",
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

os.makedirs("temp", exist_ok=True)
app.mount("/api/temp", StaticFiles(directory="temp"), name="temp")

def get_api_key() -> Optional[str]:
    """Secure extraction of the Gemini API Key."""
    return os.getenv("GEMINI_API_KEY")

# Central model config — override with GEMINI_MODEL in .env if Google retires
# this one (they retired gemini-2.5-flash for new API projects).
GEMINI_MODEL = os.getenv("GEMINI_MODEL", "gemini-3.5-flash")
# Fast model for latency-sensitive structured tasks (venue enrichment, pitches).
# flash-lite is ~5–10x faster than flash on short text/JSON. Override if quality dips.
GEMINI_FAST_MODEL = os.getenv("GEMINI_FAST_MODEL", "gemini-flash-lite-latest")
# Hard caps keep free-tier latency + cost predictable.
SCOUT_VENUE_CAP = int(os.getenv("SCOUT_VENUE_CAP", "10"))
PITCH_MAX_OUTPUT_TOKENS = int(os.getenv("PITCH_MAX_OUTPUT_TOKENS", "700"))

def get_genai_client():
    api_key = get_api_key()
    if not api_key:
         raise HTTPException(status_code=500, detail="AI features need a Gemini API key. Add GEMINI_API_KEY in backend/.env to enable them.")
    return genai.Client(api_key=api_key)


async def _generate_fast(client, *, contents, config=None, prefer_fast=True):
    """Call the fast model first; fall back to GEMINI_MODEL on hard failures.

    Pitch / scout / contract structured tasks should almost always use
    flash-lite. Full multimodal audio (Oracle) should pass prefer_fast=False.
    """
    models = [GEMINI_FAST_MODEL, GEMINI_MODEL] if prefer_fast else [GEMINI_MODEL, GEMINI_FAST_MODEL]
    last_err = None
    for model in models:
        try:
            return await client.aio.models.generate_content(
                model=model,
                contents=contents,
                config=config or {},
            )
        except Exception as err:
            last_err = err
            # Rate-limit / quota → try the other model immediately.
            msg = str(err).lower()
            if "429" in msg or "resource_exhausted" in msg or "quota" in msg:
                continue
            # Non-quota failures on the preferred model: still try the fallback once.
            continue
    raise last_err if last_err else RuntimeError("Gemini generation failed with no error detail")


def _response_text(response) -> str:
    """Best-effort plain text from a google-genai response (handles empty .text)."""
    text = getattr(response, "text", None)
    if isinstance(text, str) and text.strip():
        return text
    # Fall through candidates/parts when .text is missing or blocked.
    try:
        parts = []
        for cand in getattr(response, "candidates", None) or []:
            content = getattr(cand, "content", None)
            for part in getattr(content, "parts", None) or []:
                t = getattr(part, "text", None)
                if t:
                    parts.append(t)
        return "\n".join(parts)
    except Exception:
        return ""


def _extract_json_object(raw: str):
    """Parse JSON from model output that may include fences or trailing prose."""
    if not raw or not str(raw).strip():
        raise ValueError("empty model response")
    s = str(raw).strip()
    # Strip markdown fences ```json ... ```
    if s.startswith("```"):
        lines = s.split("\n")
        lines = lines[1:]  # drop opening fence
        if lines and lines[-1].strip().startswith("```"):
            lines = lines[:-1]
        s = "\n".join(lines).strip()
    try:
        return json.loads(s)
    except json.JSONDecodeError:
        start = s.find("{")
        end = s.rfind("}")
        if start >= 0 and end > start:
            return json.loads(s[start : end + 1])
        raise


def _normalize_oracle_payload(data: dict) -> dict:
    """Force Oracle payload into the Studio UI shape: analysis + 0-100 knobs."""
    if not isinstance(data, dict):
        raise ValueError("oracle payload is not an object")
    knobs_in = data.get("knobs") if isinstance(data.get("knobs"), dict) else {}
    # Some models nest knobs at top level.
    for key in ("sub", "air", "snap", "width"):
        if key in data and key not in knobs_in:
            knobs_in[key] = data[key]

    def clamp(v, default=50):
        try:
            n = int(round(float(v)))
        except (TypeError, ValueError):
            n = default
        return max(0, min(100, n))

    analysis = data.get("analysis") or data.get("summary") or data.get("notes") or ""
    if not isinstance(analysis, str) or not analysis.strip():
        analysis = "Mix topology scanned. Knobs set from Oracle analysis."

    return {
        "analysis": analysis.strip(),
        "knobs": {
            "sub": clamp(knobs_in.get("sub", 55)),
            "air": clamp(knobs_in.get("air", 60)),
            "snap": clamp(knobs_in.get("snap", 50)),
            "width": clamp(knobs_in.get("width", 55)),
        },
    }


def _oracle_safe_fallback(reason: str) -> dict:
    """Never leave the artist with a 500 — safe neutral-plus mastering knobs."""
    return {
        "analysis": (
            "Oracle could not fully parse the AI mix report "
            f"({reason}). Applied safe neutral-plus starting points — "
            "tweak knobs if needed, then master."
        ),
        "knobs": {"sub": 55, "air": 62, "snap": 48, "width": 58},
        "fallback": True,
        "source": "safe_fallback",
    }


def _oracle_dsp_analyze(path: str) -> dict:
    """Local spectral / dynamics analysis → Studio knobs. No AI, sub-second, always works.

    Maps band energy + crest factor + stereo correlation onto the same 0–100
    knob space the UI expects (50 = neutral).
    """
    import numpy as np

    try:
        import soundfile as sf
        audio, sr = sf.read(path, always_2d=True)
    except Exception:
        # Decode via pydub/ffmpeg when soundfile can't read the container.
        seg = AudioSegment.from_file(path)
        sr = seg.frame_rate
        samples = np.array(seg.get_array_of_samples(), dtype=np.float32)
        if seg.channels > 1:
            samples = samples.reshape((-1, seg.channels))
        else:
            samples = samples.reshape((-1, 1))
        peak = float(1 << (8 * seg.sample_width - 1))
        audio = samples / peak

    if audio.size == 0:
        return _oracle_safe_fallback("empty audio")

    # Mono mix for spectral bands; keep stereo for width.
    mono = audio.mean(axis=1).astype(np.float64)
    # Cap work on long songs (first 90s is enough for mix topology).
    max_samples = int(sr * 90)
    if mono.shape[0] > max_samples:
        mono = mono[:max_samples]
        audio = audio[:max_samples]

    # Light high-pass to ignore DC.
    mono = mono - np.mean(mono)
    n = mono.shape[0]
    # Real FFT magnitude spectrum (average of short frames would be better,
    # but a single rFFT is plenty for knob guidance and stays fast on 2GB).
    win = np.hanning(n)
    spec = np.abs(np.fft.rfft(mono * win))
    freqs = np.fft.rfftfreq(n, d=1.0 / sr)
    # Avoid log(0)
    power = np.square(spec) + 1e-12

    def band_energy(lo, hi):
        mask = (freqs >= lo) & (freqs < hi)
        if not np.any(mask):
            return 0.0
        return float(np.sum(power[mask]))

    e_sub = band_energy(20, 120)
    e_lowmid = band_energy(120, 500)
    e_mid = band_energy(500, 4000)
    e_air = band_energy(6000, min(16000, sr / 2 - 1))
    total = e_sub + e_lowmid + e_mid + e_air + 1e-12

    p_sub, p_lowmid, p_mid, p_air = e_sub / total, e_lowmid / total, e_mid / total, e_air / total

    # Crest factor (peak / RMS) → transient punch cue.
    rms = float(np.sqrt(np.mean(mono ** 2)) + 1e-12)
    peak = float(np.max(np.abs(mono)) + 1e-12)
    crest_db = 20.0 * np.log10(peak / rms)

    # Stereo width via mid/side correlation when 2ch available.
    if audio.shape[1] >= 2:
        L = audio[:, 0].astype(np.float64)
        R = audio[:, 1].astype(np.float64)
        if L.shape[0] > max_samples:
            L, R = L[:max_samples], R[:max_samples]
        corr = float(np.corrcoef(L, R)[0, 1]) if L.std() > 1e-9 and R.std() > 1e-9 else 1.0
        if np.isnan(corr):
            corr = 1.0
    else:
        corr = 1.0  # mono source

    def clamp_knob(v, lo=20, hi=80):
        # Keep suggestions in a musical range — extreme 0/100 reads as a bug in demo.
        return int(max(lo, min(hi, round(v))))

    # Knob heuristics (50 neutral):
    # thin sub → boost; muddy low-mid → cut sub slightly / leave; weak air → boost
    sub = 50 + (0.18 - p_sub) * 120          # target ~18% sub share
    if p_lowmid > 0.35:
        sub -= 6  # mud: don't pile more low end
    air = 50 + (0.12 - p_air) * 140          # target ~12% air
    # Low crest (squashed) → add snap; very spiky → ease snap
    snap = 50 + (12.0 - crest_db) * 1.8
    # High L/R correlation (narrow) → widen; already wide → ease
    width = 50 + (0.55 - corr) * 40          # corr 1.0 → ~32; corr 0.2 → ~66

    knobs = {
        "sub": clamp_knob(sub),
        "air": clamp_knob(air),
        "snap": clamp_knob(snap),
        "width": clamp_knob(width),
    }

    notes = []
    if p_sub < 0.10:
        notes.append("Sub-bass foundation is thin relative to the rest of the spectrum.")
    elif p_sub > 0.28:
        notes.append("Low end is heavy; watch for boom masking the kick/body.")
    else:
        notes.append("Low-end weight is in a workable range.")

    if p_lowmid > 0.35:
        notes.append("Low-mids are congested (boxy/mud risk around 120–500 Hz).")
    if p_air < 0.06:
        notes.append("Top end lacks air; a gentle high shelf will open the mix.")
    elif p_air > 0.22:
        notes.append("Highs are already forward — avoid harsh air boosts.")

    if crest_db < 8:
        notes.append(f"Dynamics are compressed (crest ~{crest_db:.1f} dB); snap can restore punch.")
    elif crest_db > 16:
        notes.append(f"Transients are peaky (crest ~{crest_db:.1f} dB); ease snap if it clips.")

    if corr > 0.85:
        notes.append("Stereo field is narrow/mono-leaning; width will open the sides.")
    elif corr < 0.25:
        notes.append("Image is already very wide; keep width near neutral to avoid phase haze.")

    analysis = " ".join(notes[:3]) if notes else "Mix topology measured locally. Knobs set from spectral balance."
    return {
        "analysis": analysis,
        "knobs": knobs,
        "source": "dsp",
        "metrics": {
            "band_share": {
                "sub": round(p_sub, 3),
                "lowmid": round(p_lowmid, 3),
                "mid": round(p_mid, 3),
                "air": round(p_air, 3),
            },
            "crest_db": round(crest_db, 2),
            "stereo_corr": round(corr, 3),
        },
    }

# ---------------------------------------------------------------------------
# GLOBAL LOGGING AND STARTUP
# ---------------------------------------------------------------------------

system_startup_log = []

@app.on_event("startup")
async def startup_event():
    system_startup_log.append("The Source Engine (thesourceengine.com) is starting up.")
    system_startup_log.append("Loading services…")

    # Validate FFmpeg for Audio Master Core
    try:
        subprocess.run(["ffmpeg", "-version"], capture_output=True, check=True)
        system_startup_log.append("Audio tools (FFmpeg): ready.")
    except (subprocess.CalledProcessError, FileNotFoundError):
        system_startup_log.append("Warning: FFmpeg was not found, so audio mastering won't work until it's installed.")

    # Validate Gemini SDK
    if get_api_key():
        system_startup_log.append("AI (Gemini): connected.")
    else:
        system_startup_log.append("Gemini API key is missing — AI features are off. Add GEMINI_API_KEY in backend/.env to enable them.")

    try:
        init_db()
        init_crm_tables()
        system_startup_log.append("Sign-in ready (name/email/password · sessions stored on the backend).")
        system_startup_log.append("CRM ready (leads, pitches, and activity are saved on the server).")
    except Exception as e:
        system_startup_log.append(f"Sign-in/CRM setup warning: {e}")

    if auth_required():
        system_startup_log.append("Engine features require sign-in.")
    else:
        system_startup_log.append("Sign-in is optional right now (AUTH_REQUIRED=0 — development mode).")

    system_startup_log.append(
        f"Gig data sources: Ticketmaster={'on' if ticketmaster_available() else 'off'} · "
        f"Bandsintown={'on' if bandsintown_available() else 'off'}"
    )
    system_startup_log.append("All tools ready: Studio · Radar · Legal · CRM · EPK.")
    try:
        init_profile_tables()
        system_startup_log.append("User profiles ready.")
        if os.getenv("SEED_DEMO", "").strip().lower() in ("1", "true", "yes"):
            seeded = seed_demo_artists()
            system_startup_log.append(f"Added {len(seeded)} demo artists.")
    except Exception as e:
        system_startup_log.append(f"User profile setup warning: {e}")
    try:
        system_startup_log.append(
            f"Google sign-in via Supabase: {'on' if supabase_configured() else 'off (set SUPABASE_URL + keys)'}"
        )
    except Exception:
        pass

# ---------------------------------------------------------------------------
# MODELS
# ---------------------------------------------------------------------------

class ScoutRequest(BaseModel):
    city: str
    genre: str
    tier: str
    radius: str
    timeframe: str
    country: Optional[str] = None  # ISO-2; auto-inferred from city when omitted

class NegotiateRequest(BaseModel):
    venue_offer: str

class DraftPitchRequest(BaseModel):
    venue_name: str
    venue_tier: str
    genre: str
    contact_persona: str
    payout_model: str
    artist_name: str
    agent_name: str
    agent_email: Optional[str] = None
    agent_phone: Optional[str] = None
    agent_social: Optional[str] = None
    outreach_type: str = "email" # email, call_script, dm

class AddMemoryRequest(BaseModel):
    text: str

class SearchMemoryRequest(BaseModel):
    query: str

# ---------------------------------------------------------------------------
# SYSTEM ROUTES
# ---------------------------------------------------------------------------

@app.get("/api/system-status")
async def system_status():
    stack = free_stack_manifest()
    # Match the constant used by /api/master (defined later in module load order via default).
    max_upload = int(os.getenv("MAX_MASTER_UPLOAD_BYTES", str(80 * 1024 * 1024)))
    return {
        "status": "online",
        "engine": "The Source Engine v3.0 — thesourceengine.com",
        "key_verified": bool(get_api_key()),
        "auth_required": auth_required(),
        "auth_configured": auth_configured(),
        "founding_limits": DAILY_LIMITS,
        "mastering": {
            "stack": "matchering + pedalboard + ffmpeg loudnorm",
            "max_upload_mb": max_upload // (1024 * 1024),
            "max_duration_min": 15,
            "note": "Reference match is memory-heavy; long songs may fall back to Pure mode on small hosts.",
        },
        "free_data": {
            "ticketmaster": ticketmaster_available(),
            "bandsintown": bandsintown_available(),
            "musicbrainz": True,
            "cover_art_archive": True,
            "openstreetmap": True,
            "discogs": bool(os.getenv("DISCOGS_TOKEN") or os.getenv("DISCOGS_KEY")),
            "legal_linter": True,
            "crm_sqlite": True,
        },
        "investor_stack": stack,
        "log": "\n".join(system_startup_log)
    }


class RegisterBody(BaseModel):
    name: str
    email: str
    password: str


class LoginBody(BaseModel):
    email: str
    password: str


class ForgotBody(BaseModel):
    email: str


class ResetBody(BaseModel):
    email: str
    code: str
    new_password: str


def _user_public(user: AppUser, snap: dict | None = None) -> dict:
    plan_name = (snap or {}).get("plan", {}).get("name") if snap else None
    badge = "Admin" if user.role == "admin" else (plan_name or "Member")
    if snap and snap.get("promo"):
        badge = f"{badge} · Pilot"
    return {
        "id": user.id,
        "email": user.email,
        "display_name": user.display_name,
        "role": user.role,
        "plan_id": getattr(user, "plan_id", None) or (snap or {}).get("plan_id") or "spark",
        "badge": badge,
        "auth_provider": getattr(user, "auth_provider", "password"),
    }


@app.post("/api/auth/register")
async def auth_register(body: RegisterBody):
    """Create account — name, email, password. Site opens after this."""
    user, token = register_user(body.name, body.email, body.password)
    snap = await get_usage_snapshot(user.id, user)
    return {
        "status": "success",
        "token": token,
        "user": _user_public(user, snap),
        **snap,
    }


@app.post("/api/auth/login")
async def auth_login(body: LoginBody):
    user, token = login_user(body.email, body.password)
    snap = await get_usage_snapshot(user.id, user)
    return {
        "status": "success",
        "token": token,
        "user": _user_public(user, snap),
        **snap,
    }


class GoogleAuthBody(BaseModel):
    id_token: str


@app.post("/api/auth/google")
async def auth_google(body: GoogleAuthBody):
    """
    Verify Google ID token (GIS) and issue Engine session.
    Prefer /api/auth/supabase when using Supabase Google OAuth.
    """
    client_id = (os.getenv("GOOGLE_CLIENT_ID") or "").strip()
    token = (body.id_token or "").strip()
    if not token:
        raise HTTPException(status_code=400, detail="Missing Google id_token.")

    email = name = sub = avatar = ""
    if not client_id:
        if os.getenv("GOOGLE_AUTH_MOCK", "0").strip().lower() in ("1", "true", "yes"):
            if token.startswith("mock:"):
                parts = token.split(":", 2)
                email = parts[1] if len(parts) > 1 else "pilot@example.com"
                name = parts[2] if len(parts) > 2 else "Investor Pilot"
            else:
                raise HTTPException(
                    status_code=503,
                    detail="Google Sign-In not configured. Use Supabase Google or set GOOGLE_CLIENT_ID.",
                )
        else:
            raise HTTPException(
                status_code=503,
                detail="Google Sign-In not configured. Enable Google in Supabase Auth, or set GOOGLE_CLIENT_ID.",
            )
    else:
        try:
            from google.oauth2 import id_token as google_id_token
            from google.auth.transport import requests as google_requests

            idinfo = google_id_token.verify_oauth2_token(
                token, google_requests.Request(), client_id
            )
            email = (idinfo.get("email") or "").strip().lower()
            name = (idinfo.get("name") or idinfo.get("given_name") or "").strip()
            sub = idinfo.get("sub") or ""
            avatar = idinfo.get("picture") or ""
            if not idinfo.get("email_verified", True):
                raise HTTPException(status_code=401, detail="Google email not verified.")
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(status_code=401, detail=f"Invalid Google token: {e}")

    user, session = login_or_register_google(email, name, google_sub=sub, avatar_url=avatar)
    snap = await get_usage_snapshot(user.id, user)
    profile = ensure_profile(user.id, name=user.display_name, email=user.email, avatar_url=avatar)
    return {
        "status": "success",
        "token": session,
        "user": _user_public(user, snap),
        "profile": profile,
        **snap,
    }


class SupabaseAuthBody(BaseModel):
    access_token: str


@app.post("/api/auth/supabase")
async def auth_supabase(body: SupabaseAuthBody):
    """
    Preferred path when Google is enabled in Supabase:
    Frontend: supabase.auth.signInWithOAuth({ provider: 'google' })
    Then POST the session access_token here → Engine session.
    """
    su = await fetch_supabase_user(body.access_token)
    user, session = login_or_register_google(
        su["email"],
        su.get("name") or "",
        google_sub=str(su.get("id") or ""),
        avatar_url=su.get("avatar_url") or "",
        auth_provider="supabase_google" if su.get("provider") in ("google", "supabase") else "supabase",
    )
    try:
        await optional_sync_app_user_row(su, user.plan_id or "spark")
    except Exception:
        pass
    snap = await get_usage_snapshot(user.id, user)
    profile = ensure_profile(
        user.id,
        name=user.display_name,
        email=user.email,
        avatar_url=su.get("avatar_url") or "",
    )
    return {
        "status": "success",
        "token": session,
        "user": _user_public(user, snap),
        "profile": profile,
        "auth_via": "supabase",
        **snap,
    }


@app.get("/api/auth/providers")
async def auth_providers():
    """Tell the UI which Google path is available."""
    return {
        "status": "success",
        "supabase": supabase_configured(),
        "google_gis": bool((os.getenv("GOOGLE_CLIENT_ID") or "").strip()),
        "google_mock": os.getenv("GOOGLE_AUTH_MOCK", "0").strip().lower() in ("1", "true", "yes"),
        "preferred": "supabase" if supabase_configured() else ("gis" if (os.getenv("GOOGLE_CLIENT_ID") or "").strip() else "password"),
    }


@app.post("/api/auth/logout")
async def auth_logout(request: Request):
    auth = request.headers.get("authorization") or ""
    token = auth.split(" ", 1)[1].strip() if auth.lower().startswith("bearer ") else auth.strip()
    logout_token(token)
    return {"status": "success"}


@app.post("/api/auth/forgot-password")
async def auth_forgot(body: ForgotBody):
    """
    Free password-reset: issues a 6-digit code (1h TTL).
    Without SMTP we return the code once for founding beta (document in UI).
    Always returns the same outer shape so emails can't be enumerated easily —
    but when the user exists and AUTH_SHOW_RESET_CODE is not '0', include code.
    """
    pack = create_reset_token(body.email)
    show = os.getenv("AUTH_SHOW_RESET_CODE", "1").strip().lower() not in ("0", "false", "no")
    if not pack:
        return {
            "status": "success",
            "message": "If that email is registered, a reset code was issued.",
        }
    print(f"[AUTH] Password reset code for {pack['email']}: {pack['code']}")
    out = {
        "status": "success",
        "message": "Reset code issued. Enter it with your new password.",
        "expires_in_sec": pack["expires_in_sec"],
    }
    if show:
        out["code"] = pack["code"]  # free beta: no email provider required
        out["note"] = "Beta: code returned in API (no SMTP). Set AUTH_SHOW_RESET_CODE=0 when you wire email."
    return out


@app.post("/api/auth/reset-password")
async def auth_reset(body: ResetBody):
    try:
        ok = reset_password_with_code(body.email, body.code, body.new_password)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    if not ok:
        raise HTTPException(status_code=400, detail="Invalid or expired reset code.")
    return {"status": "success", "message": "Password updated. You can sign in."}


@app.get("/api/me")
async def me(user: AppUser = Depends(require_founding_user)):
    """Profile + plan quotas + credits + promo + rich artist profile."""
    try:
        ensure_monthly_grant(user.id, user.plan_id or "spark", active_multiplier(user.id))
    except Exception:
        pass
    snap = await get_usage_snapshot(user.id, user)
    profile = ensure_profile(user.id, name=user.display_name, email=user.email)
    try:
        profile = recompute_stats_from_usage(user.id) or profile
    except Exception:
        pass
    return {
        "status": "success",
        "user": {
            **_user_public(user, snap),
            "avatar_url": profile.get("avatar_url") or user.avatar_url,
            "status": user.status,
        },
        "profile": profile,
        **snap,
    }


@app.get("/api/profile")
async def api_get_profile(user: AppUser = Depends(require_founding_user)):
    profile = ensure_profile(user.id, name=user.display_name, email=user.email)
    try:
        profile = recompute_stats_from_usage(user.id) or profile
    except Exception:
        pass
    return {"status": "success", "profile": profile}


@app.put("/api/profile")
async def api_put_profile(request: Request, user: AppUser = Depends(require_founding_user)):
    body = await request.json()
    if not isinstance(body, dict):
        raise HTTPException(status_code=400, detail="Expected JSON object.")
    existing = get_profile(user.id) or {}
    merged = {**existing, **body}
    saved = save_profile(user.id, merged)
    return {"status": "success", "profile": saved}


@app.post("/api/admin/seed-demo")
async def admin_seed_demo(user: AppUser = Depends(require_founding_user)):
    """Seed rich demo artists (Nova, Kiln, Sierra) for investor walkthrough."""
    if user.role != "admin":
        raise HTTPException(status_code=403, detail="Admin only.")
    init_profile_tables()
    artists = seed_demo_artists()
    return {
        "status": "success",
        "seeded": artists,
        "login_hint": "Password for all demo accounts: DemoArtist1!",
    }


@app.get("/api/usage")
async def usage(user: AppUser = Depends(require_founding_user)):
    snap = await get_usage_snapshot(user.id, user)
    return {"status": "success", **snap}


@app.get("/api/plans")
async def api_plans():
    """Public pricing catalog."""
    return {
        "status": "success",
        "plans": list_plans_public(),
        "credit_packs": [
            {"id": k, **v} for k, v in CREDIT_PACKS.items()
        ],
        "credit_costs": {
            "master": credit_cost("master"),
            "stems": credit_cost("stems"),
            "oracle": credit_cost("oracle"),
            "scout": credit_cost("scout"),
            "pitch": credit_cost("pitch"),
            "contract": credit_cost("contract"),
        },
    }


class PromoBody(BaseModel):
    code: str


@app.post("/api/promo/redeem")
async def api_promo_redeem(body: PromoBody, user: AppUser = Depends(require_founding_user)):
    result = redeem_promo(user.id, body.code)
    snap = await get_usage_snapshot(user.id, user)
    return {**result, **snap}


class CheckoutBody(BaseModel):
    mode: str = "subscription"  # subscription | payment
    plan_id: Optional[str] = None
    pack_id: Optional[str] = None
    interval: str = "month"


@app.post("/api/billing/checkout")
async def api_billing_checkout(body: CheckoutBody, user: AppUser = Depends(require_founding_user)):
    return create_checkout_session(
        user.id,
        user.email,
        mode=body.mode,
        plan_id=body.plan_id,
        pack_id=body.pack_id,
        interval=body.interval,
    )


@app.get("/api/billing/summary")
async def api_billing_summary(user: AppUser = Depends(require_founding_user)):
    snap = await get_usage_snapshot(user.id, user)
    return {
        "status": "success",
        "plan_id": snap.get("plan_id"),
        "plan": snap.get("plan"),
        "credits": snap.get("credits"),
        "promo": snap.get("promo"),
        "packs": [{"id": k, **v} for k, v in CREDIT_PACKS.items()],
    }


@app.post("/api/stripe/webhook")
async def api_stripe_webhook(request: Request):
    payload = await request.body()
    sig = request.headers.get("stripe-signature")
    return handle_stripe_webhook(payload, sig)


@app.get("/api/admin/users")
async def admin_users(user: AppUser = Depends(require_founding_user)):
    if user.role != "admin":
        raise HTTPException(status_code=403, detail="Admin only.")
    return {"status": "success", "users": list_users_admin(), "promo": promo_stats()}


class AdminCreditBody(BaseModel):
    user_id: str
    delta: int
    reason: str = "adjustment"


@app.post("/api/admin/credits")
async def admin_credits(body: AdminCreditBody, user: AppUser = Depends(require_founding_user)):
    if user.role != "admin":
        raise HTTPException(status_code=403, detail="Admin only.")
    bal = admin_adjust_credits(user.id, body.user_id, body.delta, body.reason)
    return {"status": "success", "balance": bal}


class AdminPlanBody(BaseModel):
    user_id: str
    plan_id: str


@app.post("/api/admin/plan")
async def admin_plan(body: AdminPlanBody, user: AppUser = Depends(require_founding_user)):
    if user.role != "admin":
        raise HTTPException(status_code=403, detail="Admin only.")
    admin_set_plan(body.user_id, body.plan_id)
    return {"status": "success", "plan_id": body.plan_id}


class AdminNoteBody(BaseModel):
    user_id: str
    body: str


@app.post("/api/admin/notes")
async def admin_notes(body: AdminNoteBody, user: AppUser = Depends(require_founding_user)):
    if user.role != "admin":
        raise HTTPException(status_code=403, detail="Admin only.")
    note = admin_add_note(user.id, body.user_id, body.body)
    return {"status": "success", "note": note}


@app.get("/api/admin/promo")
async def admin_promo(user: AppUser = Depends(require_founding_user)):
    if user.role != "admin":
        raise HTTPException(status_code=403, detail="Admin only.")
    return {"status": "success", **promo_stats()}

# ---------------------------------------------------------------------------
# PILLAR 1.5: AGENT MEMORY (Cognee Graph DB — optional subsystem)
# ---------------------------------------------------------------------------

@app.post("/api/memory/add")
async def memory_add(req: AddMemoryRequest):
    return await AgentMemory.add_document(req.text)

@app.post("/api/memory/cognify")
async def memory_cognify():
    return await AgentMemory.cognify_data()

@app.post("/api/memory/search")
async def memory_search(req: SearchMemoryRequest):
    return await AgentMemory.search_memory(req.query)

# ---------------------------------------------------------------------------
# PILLAR 2: WAR ROOM (Gig Radar Array)
# ---------------------------------------------------------------------------

def _enrich_prompt(request: ScoutRequest, real_venues: list) -> str:
    """Build a Gemini prompt that turns REAL Ticketmaster venues into full
    strategic intel, forbidding invented venue names."""
    venue_json = json.dumps(real_venues, indent=2)
    return f'''
    You are a music-booking strategist. Below is a list of REAL, VERIFIED venues
    returned by the Ticketmaster live events API for {request.city}. Each already
    has confirmed upcoming events, so they are proven-active bookers.

    REAL VENUES (do NOT invent new venues, do NOT rename these — enrich only these):
    {venue_json}

    For EACH real venue above, produce a strategic intelligence record for a
    "{request.genre}" artist targeting the "{request.tier}" tier, timeframe "{request.timeframe}".
    Use the real name, city, website_url, and upcoming_events exactly as given.
    Infer the strategic fields (payout_model, reputation, leverage, contact_persona)
    from the venue's real profile and industry norms — clearly best-guess where unknown.

    Required JSON Structure:
    {{
        "venues": [
            {{
                "name": "EXACT real venue name from the list",
                "tier": "{request.tier}",
                "contact": "Best-known booking email/phone/contact for this real venue, else the website_url",
                "contact_persona": "Name or Title of the likely talent buyer/booker",
                "contact_source": "Ticketmaster (verified) + inferred booking channel",
                "website_url": "The real website_url from the list (keep as-is), else null",
                "social_media_url": "Best-guess official social profile URL, else null",
                "similar_acts": ["Real recent/upcoming act at this venue if known", "..."],
                "payout_model": "Best-guess payout structure (Door Deal, Guarantee, Split, Unknown) for this venue tier",
                "lead_time": "Estimated booking lead time (e.g. '2 months OUT')",
                "reputation_score": "Integer 0-100 for artist fairness",
                "reputation_explanation": "1-2 sentences on WHY this score",
                "capacity": "Integer best-estimate room capacity",
                "avg_ticket_price_usd": "Integer average ticket price for this genre",
                "gross_potential_usd": "Integer (capacity * avg_ticket_price_usd)",
                "leverage_point": "One sentence pitch angle, ideally referencing a real upcoming event/gap",
                "active_search_signal": true,
                "verified_live": true
            }}
        ]
    }}
    Return strict JSON only. Include EVERY venue from the real list, no missing fields.
    '''


@app.post("/api/scout")
async def scout_gigs(request: ScoutRequest, user: AppUser = Depends(require_founding_user)):
    logs = ["Searching live ticketing listings…"]
    await assert_quota(user, "scout")
    api_key = get_api_key()
    if not api_key:
         return {"status": "error", "error": "Add your Gemini API key (GEMINI_API_KEY) in backend/.env to use gig scouting.", "log": "\n".join(logs)}

    # ---- PRIMARY: Ticketmaster · SECONDARY: Bandsintown (free public app_id) ----
    real_venues = []
    source_tag = "unknown"

    if ticketmaster_available():
        logs.append("Checking Ticketmaster for venues with active events…")
        try:
            real_venues = await fetch_ticketmaster_venues(
                city=request.city, genre=request.genre,
                radius=request.radius, timeframe=request.timeframe,
                country_code=request.country or infer_country(request.city),
            )
            source_tag = "ticketmaster_live"
            logs.append(f"Searching within country: {infer_country(request.city, request.country)}")
        except Exception as tm_err:
            logs.append(f"Couldn't get results from Ticketmaster ({str(tm_err)}).")
            real_venues = []
    else:
        logs.append("Ticketmaster API key isn't set, so skipping Ticketmaster.")

    # Merge Bandsintown venues (open public API)
    if bandsintown_available():
        logs.append("Checking Bandsintown for tours coming through your city…")
        try:
            bit = await bit_venues_near_city(request.city, request.genre)
            if bit:
                logs.append(f"Bandsintown added {len(bit)} more venues.")
                seen = {(v.get("name") or "").lower() for v in real_venues}
                for v in bit:
                    key = (v.get("name") or "").lower()
                    if key and key not in seen:
                        real_venues.append(v)
                        seen.add(key)
                source_tag = "ticketmaster+bandsintown" if source_tag == "ticketmaster_live" else "bandsintown"
        except Exception as bit_err:
            logs.append(f"Couldn't get results from Bandsintown ({bit_err}).")

    real_venues = real_venues[:SCOUT_VENUE_CAP]
    if real_venues:
        logs.append(f"Found {len(real_venues)} live venues. Adding payout estimates and booking details…")
        try:
            client = get_genai_client()
            enrich_resp = await _generate_fast(
                client,
                contents=_enrich_prompt(request, real_venues),
                config=types.GenerateContentConfig(
                    temperature=0.4,
                    response_mime_type="application/json",
                    max_output_tokens=4096,
                ),
            )
            data = json.loads(enrich_resp.text)
            if data.get("venues"):
                logs.append(f"Found {len(data['venues'])} venues booking your genre ({source_tag}).")
                await record_usage(user.id, "scout", {"city": request.city, "source": source_tag, "n": len(data["venues"])})
                try:
                    add_activity(user.id, "scout", f"Scouted {request.city} · {len(data['venues'])} venues", "radar")
                except Exception:
                    pass
                return {"status": "success", "gigs": data, "source": source_tag, "log": "\n".join(logs)}
            logs.append("Couldn't add details to those venues — trying a web search instead.")
        except Exception as enr_err:
            logs.append(f"Couldn't add venue details ({enr_err}) — trying a web search instead.")
    else:
        logs.append("No listings found on Ticketmaster or Bandsintown — searching the web instead.")


    cities_knowledge = """
    US States and Capitals Knowledge Base:
    AL: Montgomery, AK: Juneau, AZ: Phoenix, AR: Little Rock, CA: Sacramento,
    CO: Denver, CT: Hartford, DE: Dover, FL: Tallahassee, GA: Atlanta,
    HI: Honolulu, ID: Boise, IL: Springfield, IN: Indianapolis, IA: Des Moines,
    KS: Topeka, KY: Frankfort, LA: Baton Rouge, ME: Augusta, MD: Annapolis,
    MA: Boston, MI: Lansing, MN: St. Paul, MS: Jackson, MO: Jefferson City,
    MT: Helena, NE: Lincoln, NV: Carson City, NH: Concord, NJ: Trenton,
    NM: Santa Fe, NY: Albany, NC: Raleigh, ND: Bismarck, OH: Columbus,
    OK: Oklahoma City, OR: Salem, PA: Harrisburg, RI: Providence, SC: Columbia,
    SD: Pierre, TN: Nashville, TX: Austin, UT: Salt Lake City, VT: Montpelier,
    VA: Richmond, WA: Olympia, WV: Charleston, WI: Madison, WY: Cheyenne
    """
    
    prompt = f'''
    Find exactly 10 active music venues in {request.city} (within {request.radius}) that primarily book {request.genre} artists.
    The venues MUST fit the "{request.tier}" tier profile.
    Timeframe target: {request.timeframe}.
    
    You MUST search live data across:
    1. Official Booking Contacts (Emails/Webforms) on venue sites.
    2. Event history for this genre on Bandsintown/Songkick.
    3. Public social media profiles (Instagram, Twitter, Facebook) for "open booking opportunities", casting calls, or submission links if available.
    
    Required JSON Structure:
    {{
        "venues": [
            {{
                "name": "string",
                "tier": "{request.tier}",
                "contact": "string",
                "contact_persona": "Name or Title of the specific human buyer/owner",
                "contact_source": "Where the contact was found (e.g., Instagram, Official Website)",
                "website_url": "CRITICAL: The full link to their official website if found, otherwise null",
                "social_media_url": "CRITICAL: The FULL Link to their primary social media (e.g. https://instagram.com/venue) if found, otherwise null",
                "similar_acts": ["Act 1", "Act 2"],
                "payout_model": "CRITICAL: Describe the exact payout structure (Door Deal, Guarantee, Split, Unknown etc). MUST provide a best guess based on venue tier if unknown.",
                "lead_time": "CRITICAL: Estimated lead time for booking (e.g., 2 months OUT, 6 months OUT). MUST attempt a guess based on venue tier if unknown.",
                "reputation_score": "Integer 0-100 representing artist fairness",
                "reputation_explanation": "A 1-2 sentence detailed reason WHY they received this specific score (e.g. Known for late payouts, amazing sound engineer, pay-to-play packages, etc.)",
                "capacity": "Integer maximum room capacity",
                "avg_ticket_price_usd": "Integer average ticket price in USD for this genre",
                "gross_potential_usd": "Integer (capacity * avg_ticket_price_usd)",
                "leverage_point": "A single sentence identifying a weakness or opportunity to use in a pitch (e.g. 'Recently had a Thursday night cancellation' or 'Losing demographic to a rival club')",
                "active_search_signal": true
            }}
        ]
    }}
    '''
    
    try:
        client = get_genai_client()
        # Grounded search needs a tool-capable model; prefer full model, fall back to fast.
        response = await _generate_fast(
            client,
            contents=prompt,
            config=types.GenerateContentConfig(
                system_instruction=cities_knowledge,
                tools=[{"google_search": {}}],
                temperature=0.6,
                response_mime_type="application/json",
                max_output_tokens=4096,
            ),
            prefer_fast=False,
        )
        raw_text = response.text.strip()
        if raw_text.startswith("```json"):
            raw_text = raw_text[7:]
        if raw_text.startswith("```"):
            raw_text = raw_text[3:]
        if raw_text.endswith("```"):
            raw_text = raw_text[:-3]
        data = json.loads(raw_text.strip())
        
        if not data.get("venues") or len(data["venues"]) == 0:
            raise Exception("No venues found in the live search.")

        logs.append(f"Found {len(data['venues'])} venues from the web search.")
        await record_usage(user.id, "scout", {"city": request.city, "source": "grounded"})
        return {"status": "success", "gigs": data, "log": "\n".join(logs)}
        
    except Exception as fb_err:
        logs.append(f"The live search ran into a problem: {str(fb_err)}")
        # Fallback to internal knowledge
        logs.append("Falling back to well-known venues for your city…")
        try:
            fallback_prompt = f'''
            Suggest 10 legacy or famous "{request.tier}" music venues in {request.city} (within {request.radius}) for {request.genre} artists.
            Timeframe target: {request.timeframe}.
            Use internal world knowledge. Respond in strict JSON only, same schema as before (including contact_persona, contact_source, website_url, social_media_url, payout_model, lead_time, similar_acts, reputation_score, reputation_explanation, capacity, avg_ticket_price_usd, gross_potential_usd, leverage_point, and active_search_signal). Ensure NO missing fields.
            '''
            fb_response = await _generate_fast(
                client,
                contents=fallback_prompt,
                config=types.GenerateContentConfig(
                    system_instruction=cities_knowledge,
                    response_mime_type="application/json",
                    max_output_tokens=4096,
                ),
            )
            data = json.loads(fb_response.text)
            
            logs.append("Added a list of well-known venues for your city.")
            await record_usage(user.id, "scout", {"city": request.city, "source": "fallback"})
            return {"status": "success", "gigs": data, "log": "\n".join(logs)}
        except Exception as final_err:
             logs.append(f"Sorry — the search didn't work this time: {str(final_err)}")
             return {"status": "error", "error": str(final_err), "log": "\n".join(logs)}

@app.get("/api/cron/warm-gigs")
async def warm_gigs(request: Request):
    """Vercel Cron hits this every 4 hours so Ticketmaster stays warm."""
    secret = os.getenv("CRON_SECRET")
    auth = request.headers.get("authorization") or ""
    vercel_cron = request.headers.get("x-vercel-cron")
    if secret and auth != f"Bearer {secret}" and not vercel_cron:
        raise HTTPException(status_code=401, detail="cron only")

    city = "Atlanta"
    count = 0
    if ticketmaster_available():
        try:
            venues = await fetch_ticketmaster_venues(
                city=city, genre="hip hop", radius="50", timeframe="month",
                country_code="US",
            )
            count = len(venues or [])
        except Exception as err:
            return {"ok": False, "error": str(err), "city": city}
    return {"ok": True, "city": city, "venues": count}


# ---------------------------------------------------------------------------
# PILLAR 2.5: GIG RADAR (Auto-Pitch Generator)
# ---------------------------------------------------------------------------

@app.post("/api/draft-pitch")
async def draft_pitch(request: DraftPitchRequest, user: AppUser = Depends(require_founding_user)):
    logs = ["Writing your pitch…"]
    await assert_quota(user, "pitch")
    client = get_genai_client()
    
    sign_off = f"Best,\\n{request.agent_name}\\nManager for {request.artist_name}"
    if request.agent_email or request.agent_phone or request.agent_social:
        sign_off += "\\n"
        if request.agent_email:
            sign_off += f"Email: {request.agent_email}\\n"
        if request.agent_phone:
            sign_off += f"Phone: {request.agent_phone}\\n"
        if request.agent_social:
            sign_off += f"Web: {request.agent_social}\\n"

    format_instructions = ""
    max_tokens = PITCH_MAX_OUTPUT_TOKENS
    if request.outreach_type == "email":
        format_instructions = (
            f"Draft a professional but edgy email under 180 words. "
            f"Sign off exactly with:\\n{sign_off}"
        )
        max_tokens = min(max_tokens, 500)
    elif request.outreach_type == "call_script":
        format_instructions = (
            "Draft a tight 45–60 second phone script (voicemail OR live gatekeeper). "
            "Use short spoken lines and cues like [Pause] or [Enthusiastic]. "
            "Max 140 words. End with a note that you will follow up via email."
        )
        max_tokens = min(max_tokens, 450)
    elif request.outreach_type == "dm":
        format_instructions = (
            f"Draft a concise Instagram/Twitter DM under 60 words. Punchy, mobile-first. "
            f"Sign off briefly with: {request.agent_name} / {request.artist_name}."
        )
        max_tokens = min(max_tokens, 220)

    prompt = f'''
    You are '{request.agent_name}', an experienced, persuasive music manager representing the emerging {request.genre} artist "{request.artist_name}".
    Write a pitch targeting the booking buyer: "{request.contact_persona}" at the venue: "{request.venue_name}".

    CONTEXT:
    Venue Tier: {request.venue_tier}
    Expected Payout Model: {request.payout_model}

    {format_instructions}
    Return ONLY the outreach text — no preamble, no markdown fences.
    '''
    try:
        # flash-lite: pitch text is short/creative — measured ~15–18s on full flash,
        # target <4s on lite. Separate quota pool from gemini-3.5-flash free tier.
        response = await _generate_fast(
            client,
            contents=prompt,
            config=types.GenerateContentConfig(
                temperature=0.65,
                max_output_tokens=max_tokens,
            ),
        )
        logs.append("Your pitch is ready.")
        await record_usage(user.id, "pitch", {"type": request.outreach_type, "venue": request.venue_name})
        return {"status": "success", "pitch": response.text.strip(), "log": "\\n".join(logs)}
    except Exception as e:
        logs.append(f"Couldn't draft the pitch: {str(e)}")
        return {"status": "error", "error": str(e), "log": "\\n".join(logs)}

# ---------------------------------------------------------------------------
# PILLAR 3 & 4: ZION SHARK PROTOCOL (Merged)
# ---------------------------------------------------------------------------

@app.post("/api/analyze-contract")
async def analyze_contract(
    text: str = Form(None),
    file: UploadFile = File(None),
    scan_type: str = Form("contract"),
    user: AppUser = Depends(require_founding_user),
):
    logs = [f"Starting your {scan_type} review…"]
    await assert_quota(user, "contract")
    client = get_genai_client()
    
    contents = []
    extracted_text = ""
    try:
        if file:
            logs.append(f"Reading {file.filename}…")
            file_bytes = await file.read()
            
            # Determine File Type
            if file.filename.endswith('.pdf'):
                logs.append("Extracting text from your PDF…")
                reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
                for page in reader.pages:
                    extracted_text += page.extract_text() + "\n"
            elif file.filename.endswith('.docx'):
                logs.append("Extracting text from your Word document…")
                doc = Document(io.BytesIO(file_bytes))
                for para in doc.paragraphs:
                    extracted_text += para.text + "\n"
            else:
                # Fallback purely as raw text
                extracted_text = file_bytes.decode('utf-8', errors='ignore')

            contents.append(extracted_text)
        elif text:
            logs.append("Reading your pasted text…")
            contents.append(text)
        else:
            raise HTTPException(status_code=400, detail="Nothing to review yet. Paste contract text or upload a file.")

        if scan_type == 'offer':
            system_instruction = "You are an experienced, protective artist manager. Review standard venue offers critically and negotiate firmly for the artist."
            prompt = '''
            Analyze this venue booking offer. Identify lowballs or hidden fees.
            Return strictly in this JSON format:
            {
                "parties": "Venue vs Artist",
                "obligations": "What the offer details",
                "red_flags": [
                    { "clause": "Low guarantee or terrible split", "risk": "Why it's a bad deal", "fix": "Counter-offer calculation" }
                ],
                "summary": "Final negotiation stance",
                "integrity_score": 50,
                "shark_rebuttal": "Firm, professional counter-offer email to the promoter"
            }
            '''
        else:
            system_instruction = "You are an experienced music-industry contract reviewer protecting the artist. Identify predatory clauses and explain them clearly."
            prompt = '''
            Analyze this contract for an artist. Protect them.
            Return strictly in this JSON format:
            {
                "parties": "Who is involved",
                "obligations": "What the artist owes",
                "red_flags": [
                    { "clause": "Quote", "risk": "Why it's bad", "fix": "Solution" }
                ],
                "summary": "Final legal verdict",
                "integrity_score": 50,
                "shark_rebuttal": "Firm response to the lawyer"
            }
            '''

        codex_injection = """
        IMPORTANT - CODEX VOCABULARY INJECTION:
        You must actively attempt to use the following specific legal terms in your analysis 'risk', 'fix', and 'summary' fields whenever they are applicable to the document. Do not force them if they don't apply, but if they do, use these EXACT words so the UI can highlight them for the user:

        PREDATORY TERMS TO IDENTIFY: In Perpetuity, Cross-Collateralization, 360 Deal, Work For Hire, Recoupment, Leaving Member Clause, Controlled Composition, Net Profits, Black Box Royalties, Exploitation, Moral Rights, Option Periods, Right of First Refusal, Packaging Deduction, Minimum Delivery Commitment, Indemnification, Force Majeure, Territory.

        BENEFICIAL TERMS TO DEMAND/SUGGEST: Administration Deal, Co-Publishing Deal, Sync Licensing, Key Man Clause, Reversion Clause, Mutual Consent, Audit Rights, Gross Revenue, Pay or Play, Release Commitment, Escalations, Favored Nations, Sunset Clause, Cure Period.
        """

        contents.append(prompt + "\n" + codex_injection)
        
        logs.append("Analyzing the document…")
        
        # Structured legal JSON is a strong fit for flash-lite (faster + separate free-tier pool).
        response = await _generate_fast(
            client,
            contents=contents,
            config=types.GenerateContentConfig(
                system_instruction=system_instruction,
                temperature=0.2,
                response_mime_type="application/json",
                max_output_tokens=2048,
            ),
        )
        
        data = json.loads(response.text)
        # Deterministic open-source linter (always free, no AI)
        lint_source = extracted_text or text or ""
        lint = lint_text(lint_source)
        data["disclaimer"] = LEGAL_DISCLAIMER
        data["linter"] = lint
        # Prefer AI score if present; else linter hint
        if data.get("integrity_score") is None and lint.get("integrity_hint") is not None:
            data["integrity_score"] = lint["integrity_hint"]
        logs.append(
            f"{scan_type.capitalize()} review complete "
            f"({lint.get('counts', {}).get('total', 0)} terms flagged by the rule checker)."
        )
        await record_usage(user.id, "contract", {"scan_type": scan_type})
        try:
            add_activity(user.id, "scan", f"{scan_type.title()} scan · {len(data.get('red_flags') or [])} flags", "zion")
        except Exception:
            pass
        return {
            "status": "success",
            "analysis": data,
            "linter": lint,
            "disclaimer": LEGAL_DISCLAIMER,
            "log": "\n".join(logs),
        }
        
    except Exception as e:
        logs.append(f"Something went wrong during the review: {str(e)}")
        return {"status": "error", "error": str(e), "log": "\n".join(logs)}


@app.post("/api/lint-contract")
async def lint_contract_only(
    text: str = Form(None),
    file: UploadFile = File(None),
    user: AppUser = Depends(require_founding_user),
):
    """Pure free linter — no Gemini, no quota burn."""
    extracted = text or ""
    if file:
        raw = await file.read()
        name = (file.filename or "").lower()
        if name.endswith(".pdf"):
            reader = PyPDF2.PdfReader(io.BytesIO(raw))
            extracted = "\n".join((p.extract_text() or "") for p in reader.pages)
        elif name.endswith(".docx"):
            doc = Document(io.BytesIO(raw))
            extracted = "\n".join(p.text for p in doc.paragraphs)
        else:
            extracted = raw.decode("utf-8", errors="ignore")
    if not extracted.strip():
        raise HTTPException(status_code=400, detail="Paste or upload text to lint.")
    lint = lint_text(extracted)
    return {"status": "success", "linter": lint, "disclaimer": LEGAL_DISCLAIMER}

# ---------------------------------------------------------------------------
# PILLAR 5: AUDIO CORE (Matchering + Pedalboard)
# ---------------------------------------------------------------------------

def cleanup_audio_files(paths: list):
    for path in paths:
        try:
            if os.path.exists(path):
                os.remove(path)
        except Exception as e:
            print(f"Cleanup Error for {path}: {str(e)}")

# Practical cloud limits — free-tier Render OOM and laptop browsers fail above this.
MAX_MASTER_UPLOAD_BYTES = int(os.getenv("MAX_MASTER_UPLOAD_BYTES", str(80 * 1024 * 1024)))


def _validate_audio_on_disk(path: str, label: str) -> None:
    """Fail fast on empty/corrupt uploads before matchering burns CPU/RAM."""
    try:
        size = os.path.getsize(path)
    except OSError as e:
        raise HTTPException(status_code=400, detail=f"{label} could not be read: {e}")
    if size < 256:
        raise HTTPException(status_code=400, detail=f"{label} is empty or too small.")
    if size > MAX_MASTER_UPLOAD_BYTES:
        mb = MAX_MASTER_UPLOAD_BYTES // (1024 * 1024)
        raise HTTPException(
            status_code=413,
            detail=f"{label} exceeds the {mb} MB upload limit. Export a shorter bounce or lower bit depth.",
        )
    try:
        import soundfile as sf

        info = sf.info(path)
        if info.frames <= 0 or info.samplerate <= 0:
            raise HTTPException(status_code=400, detail=f"{label} has no audio frames.")
        duration = float(info.frames) / float(info.samplerate)
        if duration < 0.25:
            raise HTTPException(status_code=400, detail=f"{label} is too short to master.")
        if duration > 15 * 60:
            raise HTTPException(
                status_code=400,
                detail=f"{label} is over 15 minutes. Split the track or master a shorter section.",
            )
    except HTTPException:
        raise
    except Exception as e:
        # Non-WAV containers may still process via pedalboard/ffmpeg; only hard-fail empty
        print(f"[AUDIO CORE] Preflight probe soft-fail for {label}: {e}")


def _worker_error_detail(stdout: str, stderr: str, max_len: int = 400) -> str:
    """Surface a sanitized tail of worker logs to the client for supportability."""
    blob = (stderr or "") + "\n" + (stdout or "")
    # Prefer the last ERROR / CRITICAL line if present
    lines = [ln.strip() for ln in blob.splitlines() if ln.strip()]
    interesting = [ln for ln in lines if "ERROR" in ln.upper() or "CRITICAL" in ln.upper()]
    pick = interesting[-1] if interesting else (lines[-1] if lines else "unknown worker failure")
    pick = pick.replace("\x00", "")[:max_len]
    return pick


@app.post("/api/master")
async def master_audio(
    background_tasks: BackgroundTasks,
    target: UploadFile = File(...),
    reference: Optional[UploadFile] = File(None),
    sub: float = Form(50.0),
    air: float = Form(50.0),
    snap: float = Form(50.0),
    width: float = Form(50.0),
    warmth: float = Form(50.0),
    presence: float = Form(50.0),
    demud: float = Form(50.0),
    mono_bass: bool = Form(False),
    ref_influence: float = Form(100.0),
    lufs_target: Optional[float] = Form(None),
    output_format: str = Form("wav"),
    user: AppUser = Depends(require_founding_user),
):
    print(f"[AUDIO CORE] Ingesting mastering request. Target: {target.filename} user={user.email}")
    # Plan feature: reference matching
    if reference is not None and user.role != "admin":
        try:
            plan = effective_plan_for_user(user.id)
            if not feature_allowed(plan, "reference_master"):
                raise HTTPException(
                    status_code=403,
                    detail={
                        "error": "feature_locked",
                        "action": "reference_master",
                        "message": "Reference match mastering requires Creator or higher. Use Pure mode or upgrade.",
                        "upgrade_hint": "creator",
                    },
                )
        except HTTPException:
            raise
        except Exception:
            pass
    await assert_quota(user, "master")
    concurrent = 1
    try:
        plan = effective_plan_for_user(user.id)
        concurrent = int(plan.get("concurrent_masters") or 1)
        if user.role == "admin":
            concurrent = 99
    except Exception:
        pass
    acquire_master_slot(user.id, max_concurrent=concurrent)

    os.makedirs("temp", exist_ok=True)
    job_id = str(int(time.time()))

    # Preserve extension when possible so pedalboard/soundfile can decode mp3/flac
    def _ext(upload: UploadFile, default: str = ".wav") -> str:
        name = (upload.filename or "").lower()
        for e in (".wav", ".mp3", ".flac", ".aiff", ".aif", ".m4a", ".ogg", ".aac"):
            if name.endswith(e):
                return e
        return default

    wav_target_path = f"temp/target_{job_id}{_ext(target)}"
    wav_ref_path = f"temp/ref_{job_id}{_ext(reference) if reference else '.wav'}"
    use_ref = reference is not None and bool(getattr(reference, "filename", None))

    try:
        # Stream uploads to disk in chunks so the web server never holds the
        # full file(s) in RAM. This matters: the mastering worker (matchering)
        # peaks near the instance memory limit for full-length songs, so every
        # megabyte the parent process frees is headroom the worker needs.
        async def _save(upload, path, label: str):
            total = 0
            with open(path, "wb") as f:
                while True:
                    chunk = await upload.read(1024 * 1024)  # 1MB
                    if not chunk:
                        break
                    total += len(chunk)
                    if total > MAX_MASTER_UPLOAD_BYTES:
                        f.close()
                        try:
                            os.remove(path)
                        except OSError:
                            pass
                        mb = MAX_MASTER_UPLOAD_BYTES // (1024 * 1024)
                        raise HTTPException(
                            status_code=413,
                            detail=f"{label} exceeds the {mb} MB upload limit.",
                        )
                    f.write(chunk)

        await _save(target, wav_target_path, "Mix")
        _validate_audio_on_disk(wav_target_path, "Mix")

        # Reference is optional — no reference => "Pure mode" (DSP only, no matchering).
        if use_ref:
            await _save(reference, wav_ref_path, "Reference")
            if os.path.getsize(wav_ref_path) == 0:
                use_ref = False
            else:
                try:
                    _validate_audio_on_disk(wav_ref_path, "Reference")
                except HTTPException:
                    # Soft-drop bad reference → Pure mode rather than hard fail
                    print("[AUDIO CORE] Invalid reference — continuing in Pure mode.")
                    use_ref = False

        # Release upload buffers + force GC before the memory-heavy worker runs.
        import gc
        gc.collect()

        print(f"[DEBUG] Audio Streams pinned to disk. Mode: {'REFERENCE' if use_ref else 'PURE'}")
    except HTTPException:
        release_master_slot(user.id)
        raise
    except Exception as e:
        print("[CRITICAL] Stream Ingestion failed:")
        traceback.print_exc()
        release_master_slot(user.id)
        raise HTTPException(status_code=500, detail=f"We couldn't read your uploaded audio ({str(e)}). Please try uploading it again.")

    # Output path from worker
    mastered_wav_path = f"temp/mastered_{job_id}.wav"
    ref_arg = wav_ref_path if use_ref else "NONE"

    print("[AUDIO CORE] Dispatching to Sovereign Mastering Worker...")

    # New DSP/loudness options passed as trailing --key=value flags (backward compatible).
    flags = [
        f"--warmth={warmth}",
        f"--presence={presence}",
        f"--demud={demud}",
        f"--mono_bass={'true' if mono_bass else 'false'}",
        f"--ref_influence={ref_influence}",
    ]
    if lufs_target is not None:
        flags.append(f"--lufs={lufs_target}")

    def _run_worker(ref_path):
        # Hard ceiling so hung matchering/ffmpeg never pin the request forever.
        # Full-song reference masters on small hosts often land under 3–6 min.
        timeout_s = int(os.getenv("MASTER_WORKER_TIMEOUT_SEC", "600"))
        try:
            return subprocess.run(
                [
                    sys.executable, "matcher_worker.py",
                    wav_target_path, ref_path, mastered_wav_path,
                    str(sub), str(air), str(snap), str(width),
                    *flags,
                ],
                capture_output=True,
                text=True,
                timeout=timeout_s,
            )
        except subprocess.TimeoutExpired as te:
            out = (te.stdout or "") if isinstance(te.stdout, str) else ""
            err = (te.stderr or "") if isinstance(te.stderr, str) else f"timeout>{timeout_s}s"
            class _Fake:
                returncode = 124
                stdout = out
                stderr = err
            return _Fake()

    fell_back = False
    try:
        process = _run_worker(ref_arg)
        print(process.stdout)
        if process.returncode != 0:
            print(process.stderr)
            # Reference matching (matchering) is memory-heavy and can be killed on
            # long songs. Rather than fail the whole request, retry once in Pure
            # mode (DSP only, no matchering) so the artist still gets a master.
            if use_ref:
                print("[AUDIO CORE] Reference master failed — retrying in Pure mode (no matchering).")
                process = _run_worker("NONE")
                print(process.stdout)
                fell_back = True
            if process.returncode != 0:
                print(process.stderr)
                detail = _worker_error_detail(process.stdout or "", process.stderr or "")
                raise HTTPException(
                    status_code=500,
                    detail=(
                        "Mastering engine failed. "
                        "Try Pure mode (no reference), a shorter WAV under 80 MB, or retry after the server warms up. "
                        f"Detail: {detail}"
                    ),
                )

        # Format conversion
        final_output_path = mastered_wav_path
        if output_format.lower() in ["mp3", "flac"]:
            final_output_path = f"temp/final_{job_id}.{output_format.lower()}"
            AudioSegment.from_file(mastered_wav_path).export(final_output_path, format=output_format.lower())

        # Cleanup task
        files_to_cleanup = [wav_target_path, mastered_wav_path, final_output_path]
        if use_ref:
            files_to_cleanup.append(wav_ref_path)
        background_tasks.add_task(cleanup_audio_files, files_to_cleanup)

        media_type = "audio/wav"
        if output_format.lower() == "mp3": media_type = "audio/mpeg"
        elif output_format.lower() == "flac": media_type = "audio/flac"

        mode = "pure-fallback" if fell_back else ("reference" if use_ref else "pure")
        meters = measure_audio(mastered_wav_path)
        await record_usage(user.id, "master", {"mode": mode, "format": output_format, "lufs": meters.get("lufs_integrated")})
        try:
            add_activity(user.id, "master", f"Mastered track ({mode} · {output_format})", "audio")
        except Exception:
            pass
        # Meters as JSON headers (FileResponse body stays the audio blob)
        headers = {
            "X-Master-Mode": mode,
            "X-Master-Meters": json.dumps(meters),
            "Access-Control-Expose-Headers": "X-Master-Mode, X-Master-Meters",
        }
        return FileResponse(
            final_output_path, media_type=media_type,
            filename=f"SOURCE_MASTER.{output_format.lower()}",
            headers=headers,
        )
    finally:
        release_master_slot(user.id)

# ---------------------------------------------------------------------------
# PILLAR 6: THE ORACLE ENGINE (Mix Analysis)
# ---------------------------------------------------------------------------

@app.post("/api/oracle")
async def oracle_analysis(
    background_tasks: BackgroundTasks,
    target: UploadFile = File(...),
    user: AppUser = Depends(require_founding_user),
):
    """Oracle: local DSP first (fast + reliable), optional Gemini polish.

    Studio only needs analysis text + four knobs. DSP always supplies that in
    well under a second. Gemini may refine the write-up/knobs when it returns
    valid JSON; broken model output never 500s the request.
    """
    print(f"[ORACLE ENGINE] Intercepting unmastered payload for AI analysis: {target.filename}")
    await assert_quota(user, "oracle")
    os.makedirs("temp", exist_ok=True)
    job_id = str(uuid.uuid4())

    ext = target.filename.split('.')[-1] if '.' in target.filename else 'wav'
    temp_path = f"temp/oracle_{job_id}.{ext}"
    uploaded_file = None
    client = None

    try:
        with open(temp_path, "wb") as f:
            f.write(await target.read())

        # ---- Layer 1: deterministic DSP (always succeeds or soft-falls) ----
        try:
            dsp_oracle = _oracle_dsp_analyze(temp_path)
            print(f"[ORACLE ENGINE] DSP knobs: {dsp_oracle.get('knobs')}")
        except Exception as dsp_err:
            print(f"[ORACLE ENGINE] DSP path failed: {dsp_err}")
            traceback.print_exc()
            dsp_oracle = _oracle_safe_fallback(f"dsp: {str(dsp_err)[:80]}")

        oracle = dsp_oracle

        # ---- Layer 2: optional Gemini polish (budgeted so demos stay fast) ----
        # Default ON but hard-capped; set ORACLE_AI=0 to skip entirely.
        use_ai = os.getenv("ORACLE_AI", "1").strip() not in ("0", "false", "False", "no")
        ai_budget = float(os.getenv("ORACLE_AI_TIMEOUT_SEC", "6"))
        if use_ai and get_api_key():
            try:
                client = get_genai_client()
                print(f"[ORACLE ENGINE] Optional Gemini polish (budget {ai_budget:.0f}s)…")

                async def _ai_polish():
                    nonlocal uploaded_file
                    uploaded_file = client.files.upload(file=temp_path)
                    system_instruction = (
                        "You are The Oracle, a multi-platinum mastering engineer. "
                        "Respond with one valid JSON object only. No markdown."
                    )
                    prompt = f'''
Listen to this unmastered track. Return ONLY JSON:
{{"analysis":"2 technical sentences","knobs":{{"sub":0-100,"air":0-100,"snap":0-100,"width":0-100}}}}
50=neutral. Keep analysis under 60 words. Knob values must stay between 20 and 80.
DSP baseline knobs (you may refine): {json.dumps(dsp_oracle.get("knobs", {}))}
'''
                    # Prefer flash-lite first for latency; fall back to full model.
                    response = await _generate_fast(
                        client,
                        contents=[uploaded_file, prompt],
                        config=types.GenerateContentConfig(
                            system_instruction=system_instruction,
                            temperature=0.15,
                            response_mime_type="application/json",
                            max_output_tokens=300,
                        ),
                        prefer_fast=True,
                    )
                    raw = _response_text(response)
                    print(f"[ORACLE ENGINE] AI raw head: {raw[:160]!r}")
                    polished = _normalize_oracle_payload(_extract_json_object(raw))
                    # Re-clamp after AI so extreme values never ship.
                    for k, v in list(polished["knobs"].items()):
                        polished["knobs"][k] = int(max(20, min(80, int(v))))
                    polished["source"] = "gemini+dsp"
                    polished["dsp_knobs"] = dsp_oracle.get("knobs")
                    return polished

                oracle = await asyncio.wait_for(_ai_polish(), timeout=ai_budget)
                print("[ORACLE ENGINE] AI polish accepted.")
            except asyncio.TimeoutError:
                print(f"[ORACLE ENGINE] AI polish timed out after {ai_budget:.0f}s — DSP knobs kept.")
                oracle = dsp_oracle
                oracle["ai_error"] = f"timeout>{ai_budget:.0f}s"
            except Exception as ai_err:
                # Keep DSP result — this is the reliability win.
                print(f"[ORACLE ENGINE] AI polish skipped: {ai_err}")
                oracle = dsp_oracle
                oracle["ai_error"] = str(ai_err)[:160]
        else:
            print("[ORACLE ENGINE] AI polish disabled or no key — DSP only.")

        print("[ORACLE ENGINE] Analysis complete.")
        await record_usage(user.id, "oracle", {"source": oracle.get("source")})
        return {"status": "success", "oracle": oracle}

    except Exception as e:
        print(f"[ORACLE ENGINE] Fatal Error (soft-fail): {str(e)}")
        traceback.print_exc()
        # Soft success still consumes a unit so abuse can't free-fire oracle uploads.
        try:
            await record_usage(user.id, "oracle", {"source": "safe_fallback"})
        except Exception:
            pass
        return {"status": "success", "oracle": _oracle_safe_fallback(str(e)[:120])}
    finally:
        if uploaded_file is not None and client is not None:
            try:
                client.files.delete(name=uploaded_file.name)
            except Exception:
                pass
        background_tasks.add_task(cleanup_audio_files, [temp_path])

# ---------------------------------------------------------------------------
# PILLAR 7: OMEGA STEM EXTRACTION (Neural Separation)
# ---------------------------------------------------------------------------

@app.post("/api/extract-stems")
async def extract_stems(
    background_tasks: BackgroundTasks,
    target: UploadFile = File(...),
    user: AppUser = Depends(require_founding_user),
):
    """Open-source stems: Demucs if installed, else scipy HPSS+bands (honest labels)."""
    print(f"[STEM ENGINE] Separating stems from: {target.filename}")
    await assert_quota(user, "stems")
    os.makedirs("temp", exist_ok=True)
    job_id = str(int(time.time()))
    temp_path = f"temp/stem_source_{job_id}.wav"

    try:
        with open(temp_path, "wb") as f:
            while True:
                chunk = await target.read(1024 * 1024)
                if not chunk:
                    break
                f.write(chunk)

        result = await asyncio.to_thread(separate_stems, temp_path, "temp", job_id)
        stems_map = {
            k: f"/api/temp/{v}" for k, v in (result.get("stems") or {}).items()
        }
        background_tasks.add_task(cleanup_audio_files, [temp_path])
        await record_usage(user.id, "stems", {"method": result.get("method")})
        print(f"[STEM ENGINE] Done via {result.get('method')}")
        return {
            "status": "success",
            "stems": stems_map,
            "method": result.get("method"),
            "note": result.get("note"),
        }
    except Exception as e:
        print(f"[STEM ENGINE] Fatal Error: {str(e)}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


# ---------------------------------------------------------------------------
# FREE DATA: MusicBrainz EPK · CRM · export
# ---------------------------------------------------------------------------

class MbSearchBody(BaseModel):
    query: str
    limit: int = 8


class CrmLeadBody(BaseModel):
    id: Optional[str] = None
    venueName: Optional[str] = None
    venue_name: Optional[str] = None
    city: Optional[str] = None
    stage: Optional[str] = "scouted"
    reputationScore: Optional[int] = None
    payoutModel: Optional[str] = None
    grossPotential: Optional[float] = None
    verifiedLive: Optional[bool] = None
    meta: Optional[dict] = None


class CrmMoveBody(BaseModel):
    stage: str


class CrmPitchBody(BaseModel):
    venueName: Optional[str] = None
    venue_name: Optional[str] = None
    leadId: Optional[str] = None
    lead_id: Optional[str] = None
    outreach: Optional[str] = None
    subject: Optional[str] = None
    body: Optional[str] = None


class CrmImportBody(BaseModel):
    leads: Optional[list] = None
    pitches: Optional[list] = None
    activity: Optional[list] = None
    version: Optional[int] = 1


@app.get("/api/musicbrainz/search")
async def mb_search(q: str, user: AppUser = Depends(require_founding_user)):
    try:
        return {"status": "success", **(await search_artist(q))}
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"MusicBrainz error: {e}")


@app.get("/api/musicbrainz/releases/{artist_id}")
async def mb_releases(artist_id: str, user: AppUser = Depends(require_founding_user)):
    try:
        return {"status": "success", **(await artist_releases(artist_id))}
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"MusicBrainz error: {e}")


@app.get("/api/epk")
async def epk(q: str, user: AppUser = Depends(require_founding_user)):
    """Free EPK metadata pack: MusicBrainz + Cover Art Archive URLs."""
    try:
        pack = await epk_bundle(q)
        # Optional free Discogs layer
        try:
            d = await discogs_search(q, limit=5)
            pack["discogs"] = d
        except Exception:
            pack["discogs"] = {"results": []}
        return {"status": "success", **pack, "disclaimer": "Metadata for promotional use; verify rights before commercial reuse of images."}
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"EPK pack failed: {e}")


@app.get("/api/free/stack")
async def free_stack():
    """Public investor inventory — free live vs coming soon (no auth)."""
    return {"status": "success", **free_stack_manifest()}


@app.get("/api/free/presets")
async def free_presets(user: AppUser = Depends(require_founding_user)):
    return {"status": "success", "presets": GENRE_PRESETS}


@app.get("/api/free/release-checklist")
async def free_checklist(user: AppUser = Depends(require_founding_user)):
    return {"status": "success", "items": RELEASE_CHECKLIST}


class TourBody(BaseModel):
    cities: list[str]


@app.post("/api/free/tour-route")
async def free_tour(body: TourBody, user: AppUser = Depends(require_founding_user)):
    try:
        route = await tour_route(body.cities or [])
        return {"status": "success", **route}
    except Exception as e:
        raise HTTPException(status_code=502, detail=str(e))


@app.get("/api/free/geocode")
async def free_geocode(q: str, user: AppUser = Depends(require_founding_user)):
    hit = await nominatim_geocode(q)
    if not hit:
        return {"status": "success", "result": None}
    return {"status": "success", "result": hit}


@app.get("/api/free/discogs")
async def free_discogs(q: str, user: AppUser = Depends(require_founding_user)):
    return {"status": "success", **(await discogs_search(q))}


class OfferCompareBody(BaseModel):
    offer_a: str
    offer_b: str


@app.post("/api/free/offer-compare")
async def free_offer_compare(body: OfferCompareBody, user: AppUser = Depends(require_founding_user)):
    """Free dual-offer linter compare + optional short AI table (if Gemini on)."""
    la = lint_text(body.offer_a or "")
    lb = lint_text(body.offer_b or "")
    out = {
        "status": "success",
        "disclaimer": LEGAL_DISCLAIMER,
        "offer_a": {"linter": la, "integrity_hint": la.get("integrity_hint")},
        "offer_b": {"linter": lb, "integrity_hint": lb.get("integrity_hint")},
        "winner_hint": None,
    }
    a_s, b_s = la.get("integrity_hint") or 0, lb.get("integrity_hint") or 0
    if a_s > b_s + 5:
        out["winner_hint"] = "Offer A scores cleaner on rule-based flags (not legal advice)."
    elif b_s > a_s + 5:
        out["winner_hint"] = "Offer B scores cleaner on rule-based flags (not legal advice)."
    else:
        out["winner_hint"] = "Similar rule risk — compare money terms and exclusivity carefully."

    # Optional free-tier AI side-by-side if key present
    if get_api_key():
        try:
            client = get_genai_client()
            prompt = f"""Compare two venue/artist offers for an independent artist. Be brief.
OFFER A:
{(body.offer_a or '')[:2500]}

OFFER B:
{(body.offer_b or '')[:2500]}

Return JSON: {{
  "money_a": "one line",
  "money_b": "one line",
  "risks_a": ["..."],
  "risks_b": ["..."],
  "prefer": "A" | "B" | "mixed",
  "reason": "2 sentences"
}}"""
            resp = await _generate_fast(
                client,
                contents=prompt,
                config=types.GenerateContentConfig(
                    temperature=0.2,
                    response_mime_type="application/json",
                    max_output_tokens=700,
                ),
            )
            out["ai_compare"] = json.loads(resp.text)
        except Exception as e:
            out["ai_compare_error"] = str(e)[:160]
    return out


@app.get("/api/crm/state")
async def crm_state(user: AppUser = Depends(require_founding_user)):
    max_leads = 10
    try:
        plan = effective_plan_for_user(user.id)
        max_leads = int(plan.get("max_lead_count") or 10)
    except Exception:
        pass
    return {
        "status": "success",
        "leads": list_leads(user.id),
        "pitches": list_pitches(user.id),
        "activity": list_activity(user.id),
        "tasks": list_tasks(user.id),
        "stages": list(PIPELINE_STAGES),
        "lead_count": count_leads(user.id),
        "max_leads": max_leads if user.role != "admin" else 99999,
    }


def _assert_lead_capacity(user: AppUser, adding: int = 1) -> None:
    if user.role == "admin":
        return
    try:
        plan = effective_plan_for_user(user.id)
        max_leads = int(plan.get("max_lead_count") or 10)
    except Exception:
        max_leads = 10
    current = count_leads(user.id)
    if current + adding > max_leads:
        raise HTTPException(
            status_code=403,
            detail={
                "error": "lead_cap",
                "message": f"CRM lead cap reached ({current}/{max_leads}). Upgrade plan or remove leads.",
                "upgrade_hint": "pro",
            },
        )


@app.post("/api/crm/leads")
async def crm_upsert_lead(body: CrmLeadBody, user: AppUser = Depends(require_founding_user)):
    data = body.model_dump(exclude_none=True)
    # Only count capacity for new leads
    lid = data.get("id")
    if not lid:
        _assert_lead_capacity(user, 1)
    lead = upsert_lead(user.id, data)
    return {"status": "success", "lead": lead}


@app.post("/api/crm/leads/bulk")
async def crm_bulk(body: CrmImportBody, user: AppUser = Depends(require_founding_user)):
    incoming = body.leads or []
    _assert_lead_capacity(user, len(incoming))
    leads = bulk_upsert_leads(user.id, incoming)
    return {"status": "success", "count": len(leads), "leads": leads}


@app.patch("/api/crm/leads/{lead_id}")
async def crm_move(lead_id: str, body: CrmMoveBody, user: AppUser = Depends(require_founding_user)):
    lead = move_lead(user.id, lead_id, body.stage)
    if not lead:
        raise HTTPException(status_code=404, detail="Lead not found")
    return {"status": "success", "lead": lead}


@app.delete("/api/crm/leads/{lead_id}")
async def crm_delete(lead_id: str, user: AppUser = Depends(require_founding_user)):
    ok = delete_lead(user.id, lead_id)
    if not ok:
        raise HTTPException(status_code=404, detail="Lead not found")
    return {"status": "success"}


@app.post("/api/crm/pitches")
async def crm_pitch(body: CrmPitchBody, user: AppUser = Depends(require_founding_user)):
    pitch = add_pitch(user.id, body.model_dump(exclude_none=True))
    try:
        add_activity(user.id, "pitch", f"Pitch · {body.venueName or body.venue_name or 'venue'}", "radar")
    except Exception:
        pass
    return {"status": "success", "pitch": pitch}


@app.get("/api/crm/export")
async def crm_export(user: AppUser = Depends(require_founding_user)):
    if user.role != "admin":
        try:
            plan = effective_plan_for_user(user.id)
            if not feature_allowed(plan, "crm_export"):
                raise HTTPException(
                    status_code=403,
                    detail={
                        "error": "feature_locked",
                        "message": "CRM export requires Creator or higher.",
                        "upgrade_hint": "creator",
                    },
                )
        except HTTPException:
            raise
        except Exception:
            pass
    return {"status": "success", **export_all(user.id)}


@app.post("/api/crm/import")
async def crm_import(body: CrmImportBody, user: AppUser = Depends(require_founding_user)):
    result = import_all(user.id, body.model_dump())
    return {"status": "success", **result}


class CrmTaskBody(BaseModel):
    body: str
    lead_id: Optional[str] = None
    due_at: Optional[float] = None


@app.get("/api/crm/tasks")
async def crm_tasks(user: AppUser = Depends(require_founding_user)):
    return {"status": "success", "tasks": list_tasks(user.id, include_done=True)}


@app.post("/api/crm/tasks")
async def crm_add_task(body: CrmTaskBody, user: AppUser = Depends(require_founding_user)):
    if not (body.body or "").strip():
        raise HTTPException(status_code=400, detail="Task body required.")
    task = add_task(user.id, body.body, body.lead_id, body.due_at)
    try:
        add_activity(user.id, "task", f"Task · {body.body[:60]}", "ember")
    except Exception:
        pass
    return {"status": "success", "task": task}


@app.patch("/api/crm/tasks/{task_id}")
async def crm_task_done(task_id: str, done: bool = True, user: AppUser = Depends(require_founding_user)):
    ok = set_task_done(user.id, task_id, done)
    if not ok:
        raise HTTPException(status_code=404, detail="Task not found")
    return {"status": "success"}


if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8001, reload=True)
